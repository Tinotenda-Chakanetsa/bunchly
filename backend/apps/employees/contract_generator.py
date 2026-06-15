"""DOCX employment-contract generation.

Renders a tenant-branded contract document from an
:class:`~apps.employees.models.EmploymentContract` record + a small
override dict. Employer-side fields (legal name, address, telephone,
default benefits) come from the tenant and from the System Settings
catalogue so no business text is hard-coded — every value can be set
per tenant via ``apps.settings`` without redeploying.
"""
from __future__ import annotations

import re
from datetime import date
from decimal import Decimal
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from apps.settings.services import get_setting

from .models import ContractTemplate, Employee, EmploymentContract


# --- helpers ---------------------------------------------------------------

def _setting(tenant, key: str, fallback: str = "") -> str:
    """Read a tenant SystemSetting as a string, falling back to ``fallback``.

    Previously returned ``""`` for missing/blank values, ignoring the
    fallback — so the built-in boilerplate (working_hours_text,
    annual_leave_text, notice_text, witness_title) silently disappeared
    when the tenant hadn't filled in the System Setting.
    """
    value = get_setting(tenant, key, default=None)
    return fallback if value in (None, "") else str(value)


def _fmt_date(value) -> str:
    if not value:
        return "____________________"
    if isinstance(value, date):
        return value.strftime("%d %B %Y")
    return str(value)


def _fmt_money(amount, currency: str) -> str:
    """Format an amount as ``CCY 1,234.56`` (e.g. ``USD 5,000.00``)."""
    if amount in (None, "", 0):
        return "____________________"
    try:
        value = Decimal(str(amount))
    except (TypeError, ValueError, ArithmeticError):
        return str(amount)
    return f"{currency} {value:,.2f}"


def _employee_address(employee: Employee) -> str:
    parts = [
        employee.address_line1,
        employee.address_line2,
        employee.city,
        employee.country,
    ]
    return ", ".join(p for p in parts if p)


def _probation_months_from_contract(contract: EmploymentContract) -> int:
    """Derive probation length in whole months from the contract dates."""
    if not contract.start_date:
        return 0
    employee = contract.employee
    end = employee.probation_end_date or contract.signed_date
    if not end or end <= contract.start_date:
        return 0
    days = (end - contract.start_date).days
    return max(1, round(days / 30))


# --- main API --------------------------------------------------------------

def _tenure_paragraph(
    contract_type: str, start: str, end: str | None, probation_months: int
) -> str:
    """Pre-rendered tenure paragraph for mail-merge templates."""
    if contract_type == "full_time" or end is None:
        body = (
            f"This contract takes effect from {start} for an indefinite "
            f"period, subject to the terms set out herein."
        )
    else:
        body = (
            f"This contract takes effect from {start} to {end} and is "
            f"renewable subject to satisfactory performance. It is a "
            f"fixed-term contract and shall not create any legitimate "
            f"expectation of renewal — it automatically lapses on {end}."
        )
    if probation_months:
        plural = "" if probation_months == 1 else "s"
        body += (
            f" You shall be on {probation_months} month{plural} of probation, "
            f"during which the Employer shall assess your performance and "
            f"conduct; either party may terminate on one week's notice "
            f"during this period."
        )
    return body


def _benefits_paragraph(context: dict[str, Any]) -> str:
    """Pre-rendered benefits text reflecting the tenant's enabled blocks."""
    currency = context["currency"]
    parts: list[str] = []
    if context.get("bonus_enabled"):
        parts.append(
            "A discretionary bonus may be paid subject to availability of "
            "funds and satisfactory performance."
        )
    if context.get("transport_allowance"):
        parts.append(
            f"A monthly transport allowance of "
            f"{_fmt_money(context['transport_allowance'], currency)}."
        )
    if context.get("housing_allowance"):
        parts.append(
            f"A monthly housing allowance of "
            f"{_fmt_money(context['housing_allowance'], currency)}, payable "
            f"unless you reside in Employer-provided accommodation."
        )
    if context.get("medical_aid_enabled"):
        parts.append(
            "Eligibility to participate in the Employer's medical aid scheme."
        )
    return " ".join(parts) if parts else "No additional benefits are payable."


def build_context(
    contract: EmploymentContract, overrides: dict[str, Any] | None = None
) -> dict[str, Any]:
    """Resolve every field the template needs into a single context dict.

    Pulls employer branding from the tenant + System Settings catalogue,
    employee + position data from the related records, and lets the
    caller override anything via ``overrides`` (e.g. witness_name or
    individual paragraphs of legal text).
    """
    tenant = contract.tenant
    employee: Employee = contract.employee

    context: dict[str, Any] = {
        # --- Employer (tenant) ---------------------------------------------
        "employer_name": tenant.legal_name or tenant.name,
        "employer_address": _setting(tenant, "contract.employer_address"),
        "employer_telephone": _setting(tenant, "contract.employer_telephone"),
        "employer_country": tenant.country or "",

        # --- Employee ------------------------------------------------------
        "employee_name": employee.full_name,
        "employee_national_id": employee.national_id or "",
        "employee_address": _employee_address(employee),
        "employee_mobile": employee.phone or "",
        "employee_position": (
            contract.job_title
            or (employee.job_title.name if employee.job_title else "")
        ),
        "employee_department": (
            employee.department.name if employee.department else ""
        ),
        "employee_grade": employee.grade.name if employee.grade else "",

        # --- Contract terms ------------------------------------------------
        "contract_type": contract.contract_type,
        "contract_type_label": contract.get_contract_type_display(),
        "start_date": contract.start_date,
        "end_date": contract.end_date,
        "probation_months": _probation_months_from_contract(contract),

        # --- Remuneration --------------------------------------------------
        "basic_salary": employee.current_salary,
        "currency": employee.salary_currency or "GBP",

        # --- Defaults (configurable via System Settings) ------------------
        "working_hours_text": _setting(
            tenant, "contract.working_hours_text",
            "08:00 to 17:00 Monday to Friday, with a one-hour lunch break.",
        ),
        "annual_leave_text": _setting(
            tenant, "contract.annual_leave_text",
            "25 working days per leave year, accruing pro-rata.",
        ),
        "sick_leave_text": _setting(
            tenant, "contract.sick_leave_text",
            "Statutory sick pay in line with applicable law.",
        ),
        "notice_text": _setting(
            tenant, "contract.notice_text",
            "One (1) month written notice by either party after confirmation, "
            "or one (1) week during probation.",
        ),
        "witness_title": _setting(
            tenant, "contract.witness_title",
            "Human Resources Officer",
        ),
        "reporting_to": (
            employee.line_manager.full_name
            if employee.line_manager else "the Head of your Department"
        ),

        # --- Benefit toggles (default off unless tenant turns them on) ----
        "bonus_enabled": _setting(tenant, "contract.bonus_enabled") == "true",
        "medical_aid_enabled": _setting(
            tenant, "contract.medical_aid_enabled"
        ) == "true",
        "transport_allowance": _setting(tenant, "contract.transport_allowance"),
        "housing_allowance": _setting(tenant, "contract.housing_allowance"),
    }

    if overrides:
        # Skip empty values so a blank input in the New Contract dialog
        # falls back to the system-derived default instead of clobbering it.
        context.update(
            {k: v for k, v in overrides.items() if v not in (None, "")}
        )

    # Pre-formatted strings + composed paragraphs for mail-merge templates.
    # A tenant template can use these high-level placeholders to avoid
    # hand-coding conditional clauses in Word.
    formatted_start = _fmt_date(context["start_date"])
    formatted_end = _fmt_date(context["end_date"]) if context.get("end_date") else ""
    context.setdefault("formatted_start_date", formatted_start)
    context.setdefault("formatted_end_date", formatted_end)
    context.setdefault(
        "formatted_basic_salary",
        _fmt_money(context["basic_salary"], context["currency"]),
    )
    context.setdefault(
        "tenure_paragraph",
        _tenure_paragraph(
            context["contract_type"],
            formatted_start,
            formatted_end if context.get("end_date") else None,
            context.get("probation_months", 0),
        ),
    )
    context.setdefault("benefits_paragraph", _benefits_paragraph(context))
    return context


def render_contract(context: dict[str, Any]) -> BytesIO:
    """Render the context into a fresh ``.docx`` and return a BytesIO."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    employer = context["employer_name"]
    employee_name = context["employee_name"]
    currency = context["currency"]

    # --- Header & title ---------------------------------------------------
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(employer.upper())
    run.font.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CONTRACT OF EMPLOYMENT")
    run.font.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "This contract of employment is entered into by and between:"
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Parties ----------------------------------------------------------
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(employer).bold = True
    doc.add_paragraph('(Hereinafter referred to as "the Employer")')
    if context.get("employer_address"):
        doc.add_paragraph(f"Located at: {context['employer_address']}")
    if context.get("employer_telephone"):
        doc.add_paragraph(f"Telephone: {context['employer_telephone']}")

    doc.add_paragraph()
    p = doc.add_paragraph("AND")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(employee_name).bold = True
    doc.add_paragraph('(Hereinafter referred to as "the Employee")')
    if context.get("employee_address"):
        doc.add_paragraph(
            f"Whose residential address is: {context['employee_address']}"
        )
    if context.get("employee_national_id") or context.get("employee_mobile"):
        doc.add_paragraph(
            f"National ID / Passport No: {context['employee_national_id']}"
            f"   Mobile: {context['employee_mobile']}"
        )

    # --- 1. Position ------------------------------------------------------
    doc.add_paragraph()
    doc.add_heading("1. Position and Department", level=1)
    grade_text = (
        f", grade {context['employee_grade']}"
        if context.get("employee_grade") else ""
    )
    doc.add_paragraph(
        f"1.1 You shall be employed as a {context['contract_type_label'].lower()} "
        f"{context['employee_position']} in the "
        f"{context['employee_department']} department{grade_text}. "
        f"The Employer reserves the right to transfer you to any other "
        f"department in accordance with its operational needs."
    )

    # --- 2. Tenure --------------------------------------------------------
    doc.add_heading("2. Tenure", level=1)
    start = _fmt_date(context["start_date"])
    contract_is_fixed = (
        context["end_date"] is not None
        and context["contract_type"] != "full_time"
    )
    if contract_is_fixed:
        end = _fmt_date(context["end_date"])
        doc.add_paragraph(
            f"2.1 This contract takes effect from {start} to {end} and is "
            f"renewable subject to satisfactory performance."
        )
        doc.add_paragraph(
            "2.2 This is a fixed-term contract and shall not create any "
            f"legitimate expectation of renewal. The contract automatically "
            f"lapses on {end}."
        )
    else:
        doc.add_paragraph(
            f"2.1 This contract takes effect from {start} for an indefinite "
            f"period, subject to the terms set out herein."
        )

    if context["probation_months"]:
        doc.add_paragraph(
            f"2.3 You shall be on {context['probation_months']} month"
            f"{'s' if context['probation_months'] != 1 else ''} of "
            f"probation, during which the Employer shall assess your "
            f"performance and conduct."
        )
        doc.add_paragraph(
            "2.4 During the probation period this contract may be terminated "
            "on one week's notice by either party."
        )
    doc.add_paragraph(
        "2.5 After confirmation, termination may be effected by: "
        f"{context['notice_text']}"
    )

    # --- 3. Duties --------------------------------------------------------
    doc.add_heading("3. Duties and Responsibilities", level=1)
    doc.add_paragraph(f"3.1 You shall report directly to {context['reporting_to']}.")
    doc.add_paragraph(
        "3.2 Your responsibilities are as set out in the job description "
        "attached to this contract."
    )
    doc.add_paragraph(
        "3.3 You shall faithfully and to the best of your ability perform "
        "all duties required of your position, and shall comply with all "
        "policies, procedures and rules of the Employer."
    )

    # --- 4. Hours of work -------------------------------------------------
    doc.add_heading("4. Hours of Work and Conflict of Interest", level=1)
    doc.add_paragraph(f"4.1 Your hours of work shall be {context['working_hours_text']}")
    doc.add_paragraph(
        "4.2 Whilst employed by the Employer, you may not engage in any "
        "other employment likely to prejudice the interest of the Employer "
        "without the Employer's prior written consent."
    )

    # --- 5. Remuneration --------------------------------------------------
    doc.add_heading("5. Remuneration", level=1)
    doc.add_paragraph(
        f"5.1 Your basic salary shall be "
        f"{_fmt_money(context['basic_salary'], currency)} payable monthly "
        f"in arrears, subject to applicable statutory deductions."
    )
    doc.add_paragraph(
        "5.2 Salary increments may be made from time to time at the "
        "Employer's discretion."
    )

    # --- 6. Other benefits (each driven by tenant settings) --------------
    blocks: list[tuple[str, list[str]]] = []
    if context.get("bonus_enabled"):
        blocks.append((
            "6.1 Bonus",
            [
                "Any bonus payment shall be at the discretion of the Employer "
                "and subject to availability of funds and satisfactory "
                "performance.",
            ],
        ))
    if context.get("transport_allowance"):
        blocks.append((
            "6.2 Transport Allowance",
            [
                f"You shall receive a monthly transport allowance of "
                f"{_fmt_money(context['transport_allowance'], currency)}, "
                f"reviewable from time to time.",
            ],
        ))
    if context.get("housing_allowance"):
        blocks.append((
            "6.3 Housing Allowance",
            [
                f"You shall receive a monthly housing allowance of "
                f"{_fmt_money(context['housing_allowance'], currency)}, "
                f"unless you reside in Employer-provided accommodation.",
            ],
        ))
    if context.get("medical_aid_enabled"):
        blocks.append((
            "6.4 Medical Aid",
            [
                "You shall be eligible to participate in the Employer's "
                "medical aid scheme on the terms in force from time to time.",
            ],
        ))
    if blocks:
        doc.add_heading("6. Other Benefits", level=1)
        for heading, paragraphs in blocks:
            doc.add_paragraph(heading).runs[0].bold = True
            for text in paragraphs:
                doc.add_paragraph(text)

    # --- 7. Leave ---------------------------------------------------------
    doc.add_heading("7. Leave", level=1)
    doc.add_paragraph(f"7.1 Annual leave: {context['annual_leave_text']}")
    doc.add_paragraph(f"7.2 Sick leave: {context['sick_leave_text']}")

    # --- 8. General ------------------------------------------------------
    doc.add_heading("8. General", level=1)
    doc.add_paragraph(
        "8.1 This contract is governed by the laws of "
        f"{context.get('employer_country') or 'the jurisdiction of the Employer'}."
    )
    doc.add_paragraph(
        "8.2 This contract constitutes the entire agreement between the "
        "parties and supersedes any prior agreements or understandings."
    )

    # --- Signatures -------------------------------------------------------
    doc.add_paragraph()
    sig_title = doc.add_paragraph()
    sig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sig_title.add_run("SIGNATURE PAGE")
    run.font.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph(
        "THUS DONE AND SIGNED ON THIS ____ DAY OF ___________________ 20__"
    )
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=3)
    table.style = "Light Grid Accent 1"

    rows = [
        ("1. Employer Representative", context.get("signed_by", ""), "Date: _____________"),
        ("2. Witness", context.get("witness_name") or context["witness_title"], "Date: _____________"),
        ("3. Employee", employee_name, "Date: _____________"),
    ]
    for row, (label, who, date_cell) in zip(table.rows, rows):
        cells = row.cells
        cells[0].text = label
        cells[1].text = f"{who}\n_____________________________" if who else "_____________________________"
        cells[2].text = date_cell

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


_PLACEHOLDER_RE = re.compile(r"\{\{\s*(\w+)\s*\}\}")


def _replace_in_paragraph(paragraph, values: dict[str, str]) -> None:
    """Substitute ``{{ key }}`` tokens in a paragraph.

    Placeholders frequently span multiple runs (especially after Word
    re-saves), so we work on the paragraph's combined text and write the
    result back to the first run, clearing the rest. This loses
    fine-grained intra-paragraph formatting but is by far the most
    reliable substitution strategy.
    """
    original = paragraph.text
    if "{{" not in original:
        return
    new_text = _PLACEHOLDER_RE.sub(
        lambda m: values.get(m.group(1), m.group(0)), original
    )
    if new_text == original:
        return
    if paragraph.runs:
        for run in paragraph.runs[1:]:
            run.text = ""
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def _read_docx(template_file) -> "Document":
    """Open a Django FieldFile or path/file-like into a python-docx Document.

    FieldFile passed straight to Document() is fragile across storage
    backends — read bytes into BytesIO first.
    """
    if hasattr(template_file, "read"):
        try:
            template_file.open("rb")
        except (AttributeError, ValueError):
            pass
        try:
            template_file.seek(0)
        except (AttributeError, OSError):
            pass
        raw = template_file.read()
        try:
            template_file.close()
        except (AttributeError, ValueError):
            pass
        return Document(BytesIO(raw))
    return Document(template_file)


def parse_template_placeholders(template_file) -> list[str]:
    """Return every unique ``{{ token }}`` name used in the uploaded .docx.

    Walks body paragraphs, table cells (incl. nested), and header/footer
    regions — the same surfaces ``mail_merge_docx`` substitutes — so the
    parser sees exactly what the renderer will need to fill.
    """
    doc = _read_docx(template_file)
    found: set[str] = set()

    def collect(paragraph):
        for match in _PLACEHOLDER_RE.finditer(paragraph.text):
            found.add(match.group(1))

    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        collect(paragraph)
                    walk_tables(cell.tables)

    for paragraph in doc.paragraphs:
        collect(paragraph)
    walk_tables(doc.tables)
    for section in doc.sections:
        for region in (section.header, section.footer):
            for paragraph in region.paragraphs:
                collect(paragraph)
            walk_tables(region.tables)

    return sorted(found)


def classify_template_placeholders(tokens: list[str]) -> dict[str, list[str]]:
    """Split a token list into system-derivable vs HR-supplied buckets.

    A token is "auto" when its name matches one returned by
    :func:`available_placeholders` — those come from the tenant, employee
    and contract records. Everything else is "manual" and must be
    supplied by HR at generation time (or as an override).
    """
    catalogue = set(available_placeholders())
    auto = [t for t in tokens if t in catalogue]
    manual = [t for t in tokens if t not in catalogue]
    return {"auto": auto, "manual": manual, "all": list(tokens)}


def mail_merge_docx(template_file, context: dict[str, Any]) -> BytesIO:
    """Fill ``{{ key }}`` placeholders in a tenant-uploaded .docx template.

    Walks body paragraphs, tables (incl. nested), and header/footer
    paragraphs. Missing keys are left intact so the gap is visible in
    the output (and reported via the placeholders endpoint).
    """
    doc = _read_docx(template_file)
    values: dict[str, str] = {
        key: ("" if value is None else str(value))
        for key, value in context.items()
    }

    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, values)
                    walk_tables(cell.tables)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, values)
    walk_tables(doc.tables)
    for section in doc.sections:
        for region in (section.header, section.footer):
            for paragraph in region.paragraphs:
                _replace_in_paragraph(paragraph, values)
            walk_tables(region.tables)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def resolve_template(
    contract: EmploymentContract, template_id: str | None = None
) -> ContractTemplate | None:
    """Pick the template to use.

    Resolution order:
      1. The explicit ``template_id`` if supplied (and active).
      2. The tenant's ``is_default=True`` active template.
      3. Any active template the tenant has uploaded (most recent first) —
         so a tenant that uploaded one template but forgot the "default"
         checkbox still gets their template used, not the built-in layout.
      4. ``None`` → caller falls back to ``render_contract``.
    """
    tenant = contract.tenant
    if template_id:
        return ContractTemplate.objects.filter(
            tenant=tenant, pk=template_id, is_active=True
        ).first()
    default = ContractTemplate.objects.filter(
        tenant=tenant, is_default=True, is_active=True
    ).first()
    if default:
        return default
    return (
        ContractTemplate.objects.filter(tenant=tenant, is_active=True)
        .order_by("-created_at")
        .first()
    )


def available_placeholders() -> list[str]:
    """The list of placeholder names a tenant template can use.

    Returned by the ``placeholders`` API endpoint so HR can paste a
    reference into their template document.
    """
    return [
        # Employer
        "employer_name", "employer_address", "employer_telephone",
        "employer_country",
        # Employee
        "employee_name", "employee_national_id", "employee_address",
        "employee_mobile", "employee_position", "employee_department",
        "employee_grade",
        # Contract
        "contract_type", "contract_type_label",
        "start_date", "end_date", "formatted_start_date",
        "formatted_end_date", "probation_months",
        # Remuneration
        "basic_salary", "currency", "formatted_basic_salary",
        # Tenant-defaults paragraphs
        "working_hours_text", "annual_leave_text", "sick_leave_text",
        "notice_text", "witness_title", "reporting_to",
        "transport_allowance", "housing_allowance",
        # Pre-rendered fragments
        "tenure_paragraph", "benefits_paragraph",
    ]


def generate(
    contract: EmploymentContract,
    overrides: dict[str, Any] | None = None,
    template: ContractTemplate | None = None,
) -> BytesIO:
    """End-to-end helper.

    If a ``template`` is provided (or one is resolved as the tenant's
    default), mail-merges into that .docx. Otherwise falls back to the
    built-in ``render_contract`` layout, which is itself customisable
    via tenant ``contract.*`` SystemSettings.
    """
    context = build_context(contract, overrides)
    if template and template.template_file:
        return mail_merge_docx(template.template_file, context)
    return render_contract(context)
